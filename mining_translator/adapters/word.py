"""Word adapter for .docx files using python-docx."""

import os
from .base import BaseAdapter, TextBlock


class WordAdapter(BaseAdapter):
    supported_extensions = [".docx"]

    def extract_texts(self, filepath: str) -> list[TextBlock]:
        from docx import Document

        doc = Document(filepath)
        blocks = []

        # Extract paragraphs
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                blocks.append(TextBlock(
                    id=f"p{i}",
                    text=para.text,
                    context=f"Paragraph {i}",
                    meta={"type": "paragraph", "index": i},
                ))

        # Extract table cells
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    if cell.text.strip():
                        blocks.append(TextBlock(
                            id=f"t{ti}_r{ri}_c{ci}",
                            text=cell.text,
                            context=f"Table {ti}, row {ri}, col {ci}",
                            meta={"type": "table_cell", "table": ti, "row": ri, "col": ci},
                        ))

        return blocks

    def write_translated(self, filepath: str, blocks: list[TextBlock], output_path: str):
        from docx import Document

        doc = Document(filepath)
        lookup = {b.id: b.translated for b in blocks}

        for i, para in enumerate(doc.paragraphs):
            block_id = f"p{i}"
            if block_id in lookup:
                _replace_paragraph(para, lookup[block_id])

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    block_id = f"t{ti}_r{ri}_c{ci}"
                    if block_id in lookup:
                        _replace_cell_text(cell, lookup[block_id])

        doc.save(output_path)


def _replace_paragraph(para, new_text: str):
    """Replace paragraph text while preserving formatting of the first run."""
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.text = new_text


def _replace_cell_text(cell, new_text: str):
    """Replace cell text preserving formatting."""
    for para in cell.paragraphs:
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
            break
